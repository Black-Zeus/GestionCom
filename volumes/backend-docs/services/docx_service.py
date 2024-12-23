from docx import Document
import os
from config.minio_config import get_minio_client, MINIO_BUCKET_TEMP, TMP_OUTPUT

minio_client = get_minio_client()

async def generate_docx():
    docx_path = f"{TMP_OUTPUT}/docx_demo.docx"
    doc = Document()
    doc.add_heading("Reporte Word", 0)
    doc.add_paragraph("Este es un documento Word generado correctamente.")
    doc.save(docx_path)

    # Subir a MinIO (Bucket de backups)
    minio_client.fput_object(MINIO_BUCKET_TEMP, "xlsx_demo.xlsx", docx_path)
    os.remove(docx_path)

    return {"message": "Documento Word generado y subido a MinIO"}
